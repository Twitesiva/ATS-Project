"""Extract plain text from PDF and DOCX resume files."""
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
from backend.config import UPLOAD_FOLDER

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    from docx import Document
except Exception:
    Document = None


# PERFORMANCE OPTIMIZATION – NON-BREAKING: Text preprocessing to reduce embedding computation
def _preprocess_text(text):
    """
    Clean and preprocess text before embedding.
    Removes excessive whitespace while preserving semantic content.
    """
    if not text:
        return ""
    
    # Remove excessive newlines and whitespace
    import re
    # Replace multiple newlines with single newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Replace multiple spaces with single space
    text = re.sub(r' {2,}', ' ', text)
    # Remove common footer/header patterns (page numbers, etc.)
    text = re.sub(r'\n\s*Page\s+\d+\s*(of\s*\d+)?\s*\n', '\n', text, flags=re.IGNORECASE)
    # Remove email/phone signature blocks (common patterns)
    text = re.sub(r'\n[_-]{10,}\n.*?(?:email|phone|contact).*', '', text, flags=re.IGNORECASE | re.DOTALL)
    
    return text.strip()


def _read_pdf(file_path):
    """Read PDF text with pdfplumber (primary) and PyPDF2 (fallback)."""
    text_parts = []
    
    # Try pdfplumber first (better layout preservation)
    try:
        if pdfplumber is None:
            raise ImportError("pdfplumber not available")
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        if text_parts:
            raw_text = "\n".join(text_parts)
            return _preprocess_text(raw_text)
    except Exception as e:
        print(f"[PARSER] pdfplumber failed: {e}. Falling back to PyPDF2.")
    
    # Fallback to PyPDF2
    try:
        if PyPDF2 is None:
            raise ImportError("PyPDF2 not available")
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        raw_text = "\n".join(text_parts) if text_parts else ""
        return _preprocess_text(raw_text)
    except Exception as e:
        print(f"[PARSER] PyPDF2 also failed: {e}")
        return ""


def _read_docx(file_path):
    if Document is None:
        return ""
    doc = Document(file_path)
    chunks = []
    chunks.extend(p.text.strip() for p in doc.paragraphs if p.text and p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    chunks.append(cell_text)
    if not chunks:
        try:
            import zipfile
            import xml.etree.ElementTree as ET

            with zipfile.ZipFile(file_path) as zf:
                xml_parts = [name for name in zf.namelist() if name.startswith("word/") and name.endswith(".xml")]
                for part in xml_parts:
                    xml_bytes = zf.read(part)
                    root = ET.fromstring(xml_bytes)
                    for node in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                        value = (node.text or "").strip()
                        if value:
                            chunks.append(value)
        except Exception:
            pass
    raw_text = "\n".join(chunks)
    return _preprocess_text(raw_text)


def parse_resume_file(full_path, original_name):
    """Parse one file; return dict with path (relative), original_name, text."""
    ext = os.path.splitext(full_path)[-1].lower()
    if ext == ".pdf":
        text = _read_pdf(full_path)
    elif ext in (".docx", ".doc"):
        text = _read_docx(full_path)
    else:
        text = ""
    return {"path": os.path.basename(full_path), "original_name": original_name, "text": (text or "").strip()}


# PERFORMANCE OPTIMIZATION – NON-BREAKING: Parallel file parsing using thread pool
def _parse_single_resume(item):
    """Helper function for parallel parsing."""
    path = item.get("path") or item.get("original_name", "")
    original_name = item.get("original_name") or path
    full_path = os.path.join(UPLOAD_FOLDER, path)
    
    if not os.path.isfile(full_path):
        return {"path": path, "original_name": original_name, "text": ""}
    
    return parse_resume_file(full_path, original_name)


def parse_resumes_from_paths(resume_paths, max_workers=4):
    """
    resume_paths: list of { path, original_name } (path is relative name under UPLOAD_FOLDER).
    Return list of { path, original_name, text }.
    
    PERFORMANCE OPTIMIZATION – NON-BREAKING: Uses parallel processing for multiple files.
    """
    if not resume_paths:
        return []
    
    # For single file, skip thread overhead
    if len(resume_paths) == 1:
        return [_parse_single_resume(resume_paths[0])]
    
    # For multiple files, use thread pool for parallel I/O
    results = []
    with ThreadPoolExecutor(max_workers=min(max_workers, len(resume_paths))) as executor:
        # Submit all tasks
        future_to_item = {
            executor.submit(_parse_single_resume, item): item 
            for item in resume_paths
        }
        
        # Collect results as they complete
        for future in as_completed(future_to_item):
            try:
                result = future.result()
                results.append(result)
            except Exception as e:
                # On error, return empty text for that file
                item = future_to_item[future]
                path = item.get("path") or item.get("original_name", "")
                original_name = item.get("original_name") or path
                results.append({"path": path, "original_name": original_name, "text": ""})
    
    return results
